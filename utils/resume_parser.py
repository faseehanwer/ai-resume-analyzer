"""
Resume Parser Utility
Extracts raw text from PDF, DOCX, and TXT resume files.
"""

import re
from io import BytesIO


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
        doc = docx.Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")


def clean_text(text: str) -> str:
    """Normalize extracted text — collapse whitespace, fix line breaks."""
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_resume(file) -> dict:
    """
    Parse an uploaded resume file and return extracted metadata.

    Args:
        file: A Streamlit UploadedFile object.

    Returns:
        dict with keys: raw_text, word_count, char_count, file_name, file_type
    """
    file_bytes = file.read()
    file_name = file.name.lower()

    if file_name.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
        file_type = "PDF"
    elif file_name.endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
        file_type = "DOCX"
    elif file_name.endswith(".txt"):
        raw_text = file_bytes.decode("utf-8", errors="ignore")
        file_type = "TXT"
    else:
        raise ValueError(f"Unsupported file type: {file.name}")

    raw_text = clean_text(raw_text)

    if not raw_text:
        raise ValueError("Could not extract any text from the uploaded file.")

    return {
        "raw_text": raw_text,
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
        "file_name": file.name,
        "file_type": file_type,
    }
